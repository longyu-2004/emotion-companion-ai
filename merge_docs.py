import os, re, copy
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = "/sessions/gallant-eloquent-hypatia/mnt/大三下软件项目管理作业/情感陪护"
DOCS = [
    (1, "实践一  项目建议书", f"{BASE}/1-2/第六组-基于AI大语言模型的情感陪护虚拟数字人系统-项目建议书.docx"),
    (2, "实践二  项目章程", f"{BASE}/1-2/第六组-基于AI大语言模型的情感陪护虚拟数字人系统-项目章程.docx"),
    (3, "实践三  项目生存期模型确定报告", f"{BASE}/1-2/第六组-基于AI大语言模型的情感陪护虚拟数字人系统-项目生存期模型确定报告.docx"),
    (4, "实践四  项目需求管理", f"{BASE}/3-5/龙宇-基于AI大语言模型的情感陪护虚拟数字人系统/第六组-基于AI大语言模型的情感陪护虚拟数字人系统-项目需求管理.docx"),
    (5, "实践五  WBS", f"{BASE}/3-5/龙宇-基于AI大语言模型的情感陪护虚拟数字人系统/第六组-基于AI大语言模型的情感陪护虚拟数字人系统-WBS.docx"),
    (6, "实践六  成本估算", f"{BASE}/3-5/龙宇-基于AI大语言模型的情感陪护虚拟数字人系统/第六组-基于AI大语言模型的情感陪护虚拟数字人系统-成本估算.docx"),
    (7, "实践七  项目进度管理", f"{BASE}/6-8/第六组-基于AI大语言模型的情感陪护虚拟数字人系统-项目进度管理(实践六).docx"),
    (8, "实践八  项目进度编排报告", f"{BASE}/6-8/第六组-基于AI大语言模型的情感陪护虚拟数字人系统-实践七_项目进度编排报告.docx"),
    (9, "实践九  配置管理计划", f"{BASE}/6-8/第六组-基于AI大语言模型的情感陪护虚拟数字人系统-配置管理计划.docx"),
    (10, "实践十  人力计划与沟通计划", f"{BASE}/9-12/第六组-基于AI大语言模型的情感陪护虚拟数字人系统-实践九_人力计划与沟通计划.docx"),
    (11, "实践十一  风险计划", f"{BASE}/9-12/第六组-基于AI大语言模型的情感陪护虚拟数字人系统-实践十_风险计划(1).docx"),
    (12, "实践十二  系统原型的设计实现", f"{BASE}/9-12/第六组-基于AI大语言模型的情感陪护虚拟数字人系统-实践十一-系统原型的设计实现.docx"),
]
OUTPUT = os.path.join(BASE, "第六组-软件项目管理课程实践报告-完整版v2.docx")

F_TITLE = "黑体"
F_BODY = "宋体"
F_EN = "Times New Roman"
SZ_H1 = Pt(16)
SZ_H2 = Pt(14)
SZ_H3 = Pt(12)
SZ_BODY = Pt(12)
SZ_SMALL = Pt(10.5)
CLR = RGBColor(0, 0, 0)
LSP = Pt(22)

SKIP = [
    r'^项目名称[：:]', r'^项目单位[：:]', r'^团队名称[：:]', r'^团\s*队[：:]',
    r'^团队成员[：:]', r'^组\s*员[：:]', r'^指导教师[：:]', r'^主要用户群体[：:]',
    r'^提交日期[：:]', r'^编制日期[：:]', r'^开发周期[：:]', r'^生存期模型[：:]',
    r'^实践目的[：:]，', r'^专业[：:]', r'^预算[：:]', r'^项目组别[：:]',
    r'^【项目名称】', r'^【团队名称】', r'^【团队成员】', r'^【开发周期】',
    r'^基于AI大语言模型.*建设项目建议书', r'^基于AI大语言模型.*项目章程$',
    r'^项目生存期模型确定报告$', r'^软件项目管理实践报告$',
    r'^软件项目管理实践.*报告$', r'^项目需求管理$',
    r'^项目范围管理与WBS实践$', r'^软件项目成本估算报告$',
    r'^配置管理计划$', r'^实验[五六七八九十]$', r'^目\s*录$',
    r'^基于AI大语言模型的情感陪护虚拟数字人系统$',
    r'^实践[一二三四五六七八九十]+[：:\s]',
]

def should_skip(text):
    t = text.strip()
    if not t or len(t) < 2:
        return True
    for p in SKIP:
        if re.match(p, t):
            return True
    return False

def srf(run, cn=F_BODY, sz=SZ_BODY, bold=False):
    run.font.name = F_EN
    run.font.size = sz
    run.font.bold = bold
    run.font.color.rgb = CLR
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = parse_xml('<w:rFonts %s />' % nsdecls('w'))
        rPr.insert(0, rf)
    rf.set(qn('w:eastAsia'), cn)
    rf.set(qn('w:ascii'), F_EN)
    rf.set(qn('w:hAnsi'), F_EN)

def spf(para, al=WD_ALIGN_PARAGRAPH.JUSTIFY, ind=None, ls=LSP, bef=Pt(0), aft=Pt(0)):
    f = para.paragraph_format
    f.alignment = al
    f.line_spacing = ls
    f.line_spacing_rule = 4
    f.space_before = bef
    f.space_after = aft
    if ind is not None:
        f.first_line_indent = ind

def dh(text):
    t = text.strip()
    if not t:
        return 0
    if re.match(r'^[一二三四五六七八九十]+[、．.]', t) or re.match(r'^第[一二三四五六七八九十]+部分', t):
        return 2
    if re.match(r'^\d+[\.\、]', t):
        return 3
    return 0

def has_img(elem):
    return len(elem.findall('.//' + qn('w:drawing'))) > 0 or len(elem.findall('.//' + qn('w:pict'))) > 0

def get_rids(elem):
    rids = set()
    for b in elem.iter(qn('a:blip')):
        r = b.get(qn('r:embed'))
        if r: rids.add(r)
    for i in elem.iter(qn('v:imagedata')):
        r = i.get(qn('r:id'))
        if r: rids.add(r)
    return rids

def copy_imgs(src, dst, elem):
    for rid in get_rids(elem):
        try:
            rel = src.part.rels.get(rid)
            if not rel: continue
            blob = rel.target_part.blob
            ct = rel.target_part.content_type
            nrid = dst.part.relate_to(BytesIO(blob), ct)
            for blip in elem.iter(qn('a:blip')):
                if blip.get(qn('r:embed')) == rid:
                    blip.set(qn('r:embed'), nrid)
            for img in elem.iter(qn('v:imagedata')):
                if img.get(qn('r:id')) == rid:
                    img.set(qn('r:id'), nrid)
        except: pass

def fmt_p(para, h=0):
    t = para.text.strip()
    if h == 2:
        for r in para.runs: srf(r, cn=F_TITLE, sz=SZ_H2, bold=True)
        spf(para, al=WD_ALIGN_PARAGRAPH.LEFT, bef=Pt(12), aft=Pt(6))
    elif h == 3:
        for r in para.runs: srf(r, cn=F_TITLE, sz=SZ_H3, bold=True)
        spf(para, al=WD_ALIGN_PARAGRAPH.LEFT, bef=Pt(6), aft=Pt(3))
    else:
        is_ref = t.startswith("[") or re.match(r'^\[\d+\]', t)
        sz = SZ_SMALL if is_ref else SZ_BODY
        for r in para.runs: srf(r, sz=sz)
        spf(para, ind=Cm(0.74))

def fmt_tbl(elem):
    for r in elem.iter(qn('w:r')):
        rp = r.find(qn('w:rPr'))
        if rp is None:
            rp = parse_xml('<w:rPr %s/>' % nsdecls('w'))
            r.insert(0, rp)
        for tag, val in [('w:sz', '21'), ('w:szCs', '21'), ('w:color', '000000')]:
            e = rp.find(qn(tag))
            if e is None:
                e = parse_xml('<%s %s/>' % (tag, nsdecls('w')))
                rp.append(e)
            e.set(qn('w:val'), val)
        rf = rp.find(qn('w:rFonts'))
        if rf is None:
            rf = parse_xml('<w:rFonts %s/>' % nsdecls('w'))
            rp.insert(0, rf)
        rf.set(qn('w:eastAsia'), F_BODY)
        rf.set(qn('w:ascii'), F_EN)
        rf.set(qn('w:hAnsi'), F_EN)

def main():
    print("=" * 60)
    print("  Merge 12 docs v2")
    print("=" * 60)
    m = Document()
    sec = m.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)
    st = m.styles['Normal']
    st.font.name = F_EN
    st.font.size = SZ_BODY
    st.font.color.rgb = CLR
    st.paragraph_format.line_spacing = LSP
    st.paragraph_format.line_spacing_rule = 4
    rp = st.element.get_or_add_rPr()
    rf = rp.find(qn('w:rFonts'))
    if rf is None:
        rf = parse_xml('<w:rFonts %s />' % nsdecls('w'))
        rp.insert(0, rf)
    rf.set(qn('w:eastAsia'), F_BODY)
    for hn, hs in [('Heading 1', SZ_H1), ('Heading 2', SZ_H2), ('Heading 3', SZ_H3)]:
        try:
            s = m.styles[hn]
            s.font.name = F_EN; s.font.size = hs; s.font.bold = True; s.font.color.rgb = CLR
            srp = s.element.get_or_add_rPr()
            srf2 = srp.find(qn('w:rFonts'))
            if srf2 is None:
                srf2 = parse_xml('<w:rFonts %s />' % nsdecls('w'))
                srp.insert(0, srf2)
            srf2.set(qn('w:eastAsia'), F_TITLE)
        except: pass
    print("\n  Cover...")
    for _ in range(6): m.add_paragraph()
    p = m.add_paragraph(); r = p.add_run("软件项目管理课程实践报告")
    srf(r, cn=F_TITLE, sz=Pt(26), bold=True); spf(p, al=WD_ALIGN_PARAGRAPH.CENTER, aft=Pt(24))
    m.add_paragraph()
    p = m.add_paragraph(); r = p.add_run("基于AI大语言模型的情感陪护虚拟数字人系统")
    srf(r, cn=F_TITLE, sz=Pt(18), bold=True); spf(p, al=WD_ALIGN_PARAGRAPH.CENTER, aft=Pt(48))
    for item in ["项目名称：基于AI大语言模型的情感陪护虚拟数字人系统", "项目组别：第六组", "组    员：龙宇、姜乐、姜怡琳", "指导教师：刘志远"]:
        p = m.add_paragraph(); r = p.add_run(item); srf(r, sz=Pt(14)); spf(p, al=WD_ALIGN_PARAGRAPH.CENTER, ls=Pt(28))
    m.add_paragraph()
    p = m.add_paragraph(); r = p.add_run("2026年6月"); srf(r, sz=Pt(14)); spf(p, al=WD_ALIGN_PARAGRAPH.CENTER)
    m.add_page_break()
    p = m.add_paragraph(); r = p.add_run("目    录")
    srf(r, cn=F_TITLE, sz=Pt(16), bold=True); spf(p, al=WD_ALIGN_PARAGRAPH.CENTER, bef=Pt(24), aft=Pt(24))
    tp = m.add_paragraph()
    r1 = tp.add_run(); r1._element.append(parse_xml('<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w')))
    r2 = tp.add_run(); r2._element.append(parse_xml('<w:instrText %s xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>' % nsdecls('w')))
    r3 = tp.add_run(); r3._element.append(parse_xml('<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w')))
    r4 = tp.add_run("请在Word中右键此处选择「更新域」生成目录"); srf(r4, sz=Pt(12)); r4.font.color.rgb = RGBColor(128, 128, 128)
    r5 = tp.add_run(); r5._element.append(parse_xml('<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w')))
    for pn, pt, dp in DOCS:
        print(f"\n  [{pn:2d}/12] {pt}")
        if not os.path.exists(dp): print("         SKIP: not found"); continue
        try: src = Document(dp)
        except Exception as e: print(f"         ERROR: {e}"); continue
        m.add_page_break()
        h1 = m.add_paragraph(); h1.style = m.styles['Heading 1']
        hr = h1.add_run(pt); srf(hr, cn=F_TITLE, sz=SZ_H1, bold=True)
        spf(h1, al=WD_ALIGN_PARAGRAPH.CENTER, bef=Pt(12), aft=Pt(12))
        sk = cp = ic = 0
        for el in src.element.body:
            tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
            if tag == 'p':
                txt = ""
                for t in el.iter(qn('w:t')): txt += (t.text or "")
                if should_skip(txt): sk += 1; continue
                hd = dh(txt)
                try:
                    ec = copy.deepcopy(el)
                    if has_img(el): copy_imgs(src, m, ec); ic += 1
                    np = m.add_paragraph(); np._element.clear()
                    for ch in ec: np._element.append(ch)
                    if not list(np._element.findall(qn('w:r'))) and txt:
                        re2 = parse_xml('<w:r %s><w:rPr><w:rFonts w:eastAsia="%s" w:ascii="%s" w:hAnsi="%s"/><w:sz w:val="24"/><w:color w:val="000000"/></w:rPr><w:t xml:space="preserve">%s</w:t></w:r>' % (nsdecls('w'), F_BODY, F_EN, F_EN, txt))
                        np._element.append(re2)
                    fmt_p(np, hd)
                    if hd >= 2: np.style = m.styles['Heading 2' if hd == 2 else 'Heading 3']
                    cp += 1
                except: pass
            elif tag == 'tbl':
                try:
                    sp = m.add_paragraph(); spf(sp, ls=Pt(6), bef=Pt(6), aft=Pt(6))
                    tc = copy.deepcopy(el); fmt_tbl(tc); m.element.body.append(tc); cp += 1
                except: pass
        print(f"         copied={cp} skipped={sk} imgs={ic}")
    print(f"\n  Saving: {OUTPUT}")
    m.save(OUTPUT)
    print(f"  Size: {os.path.getsize(OUTPUT) / 1024:.1f} KB")
    print("=" * 60)
    print("  Done! Right-click TOC in Word -> Update Field")
    print("=" * 60)

if __name__ == "__main__":
    main()
